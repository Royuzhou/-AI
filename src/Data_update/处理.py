import re
from pathlib import Path
from docx import Document
from sentence_transformers import SentenceTransformer
import concurrent.futures
import os
from pinecone import Pinecone, ServerlessSpec
import hashlib
import unicodedata


# 初始化嵌入模型
embedding_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

# 初始化Pinecone客户端
pc = Pinecone(
    api_key='pcsk_6f3yXX_CK5s2GTwXKPQVectvkzKqdQvBArp4Q6FnHhxYV5UDGiVBaFTEBN9AiErWv58Kd2'
)

index = pc.Index("software")


def split_legal_clauses(text):
    """按条款编号分割文本，保留完整条款内容"""
    text = str(text)
    # 改进的正则表达式，增加前后边界判断
    pattern = r'(?<!\S)(第[一二三四五六七八九十零百千万]+(?:\.\d+)?[条款项])(?=\s|$)'
    matches = list(re.finditer(pattern, text))
    
    clauses = []
    prev_end = 0
    
    for i, match in enumerate(matches):
        clause_num = match.group(1)
        start = match.end()
        
        # 确定当前条款结束位置
        if i < len(matches)-1:
            next_start = matches[i+1].start()
        else:
            next_start = len(text)
        
        # 提取条款内容（包含编号后的标点）
        clause_content = text[start:next_start].strip()

        # 处理前导内容（标题/说明）
        if prev_end < start:
            pre_content = text[prev_end:start].strip()
            if pre_content:
                clauses.append({
                    "clause_num": f"前导_{pre_content}",
                    "content": pre_content,
                    "metadata": {"text": pre_content}
                })
        
        clauses.append({
            "clause_num": clause_num,
            "content": clause_content,
            "metadata": {"text": clause_num}
        })
        prev_end = next_start
    
    return clauses

def process_docx(file_path):
    """增强版合同解析"""
    try:
        doc = Document(file_path)
        file_name = Path(file_path).stem.replace(' ', '_')
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        # 改进的分割逻辑
        clauses = split_legal_clauses(full_text)
        clauses = [c for c in clauses if c["content"]]
        
        processed_chunks = []
        for clause in clauses:
            content = clause["content"]
            # 控制内容长度并生成元数据
            content_part = content[:512] if len(content) > 512 else content
            processed_chunks.append({
                "clause_num": clause["clause_num"],
                "content": content_part,
                "metadata": {
                    "file_id": file_name,
                    "clause_num": clause["clause_num"],
                    "original_text": content_part
                }
            })
        return processed_chunks
    except Exception as e:
        print(f"处理文件 {file_name} 失败: {str(e)}")
        return []

def process_docx(file_path):
    """修复文件名获取方式"""
    try:
        # 直接使用路径字符串处理DOCX
        doc = Document(file_path)
        # 正确获取文件名（不依赖对象属性）
        file_name = os.path.basename(file_path)
        full_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
        clauses = split_legal_clauses(full_text)
        clauses = [c for c in clauses if c["content"]]
        
        processed_chunks = []
        for clause in clauses:
            content = clause["content"]
            # 控制内容长度并生成元数据
            content_part = content[:512] if len(content) > 512 else content
            metadata = {
                "file_id": file_name,  # 使用文件名而非对象属性
                "clause_num": clause["clause_num"],
                "original_text": content_part
            }
            processed_chunks.append({
                "clause_num": clause["clause_num"],
                "content": content_part,
                "metadata": metadata
            })
        return processed_chunks
    except Exception as e:
        print(f"处理文件 {file_name} 失败: {str(e)}")
        return []

def generate_vector_id(file_id, clause_num):
    """生成纯ASCII向量ID"""
    # 清除非ASCII字符
    clean_file_id = unicodedata.normalize('NFKD', file_id).encode('ASCII', 'ignore').decode()
    clean_clause_num = unicodedata.normalize('NFKD', clause_num).encode('ASCII', 'ignore').decode()
    
    # 生成哈希值
    hash_digest = hashlib.md5(clean_clause_num.encode()).hexdigest()
    short_hash = hash_digest[:8]
    
    # 组合ASCII ID
    return f"{clean_file_id}_{clean_clause_num}_{short_hash}"

def index_to_pinecone(file_id, chunks):
    """批量存入Pinecone"""
    vectors = []
    for chunk in chunks:
        try:
            vector = embedding_model.encode(chunk["content"]).tolist()
            vector_id = generate_vector_id(file_id, chunk["clause_num"])
            vectors.append((vector_id, vector, chunk["metadata"]))
        except Exception as e:
            print(f"向量生成失败: {str(e)}")
            continue
    
    # 分批次上传（每批≤2MB）
    batch_size = 25
    for i in range(0, len(vectors), batch_size):
        batch = vectors[i:i+batch_size]
        index.upsert(vectors=batch)

def process_folder(folder_path):
    """主函数：遍历文件夹→多线程处理→存储"""
    doc_files = list(Path(folder_path).rglob("*.docx"))
    print(f"发现 {len(doc_files)} 个DOCX文件")

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = {
            executor.submit(process_docx, str(file)): file.name
            for file in doc_files
        }
        
        for future in concurrent.futures.as_completed(futures):
            file_name = futures[future]
            try:
                chunks = future.result()
                if chunks:
                    index_to_pinecone(file_name, chunks)
            except Exception as e:
                print(f"文件 {file_name} 处理失败: {str(e)}")

if __name__ == "__main__":
    process_folder("C:\\Users\\Roy\\Desktop\\商业合同")