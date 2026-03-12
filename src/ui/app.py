"""
Flask 后端服务 - 合同修订智能体 UI
提供 API 接口和前端页面
"""

import os
import sys
import json
import re
from flask import Flask, render_template, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from src.agent import ContractRevisionAgent
from config import CONFIG

app = Flask(__name__)
CORS(app)

app.config['JSON_AS_ASCII'] = False

agent = None


def init_agent():
    """初始化智能体"""
    global agent
    if agent is None:
        try:
            agent = ContractRevisionAgent(CONFIG)
            print("✓ 智能体初始化成功")
        except Exception as e:
            print(f"✗ 智能体初始化失败: {e}")
            agent = None


def parse_contract_result(result_text):
    """解析合同修订结果，提取修订后的合同和修改建议"""
    
    revised_contract = ""
    suggestions = []
    references = []
    
    lines = result_text.split('\n')
    
    current_section = None
    in_revised_contract = False
    in_suggestions = False
    current_suggestion = None
    
    for line in lines:
        line_stripped = line.strip()
        
        if '【修订后的完整合同】' in line_stripped:
            in_revised_contract = True
            in_suggestions = False
            continue
        
        if '【修改建议】' in line_stripped:
            in_revised_contract = False
            in_suggestions = True
            continue
        
        if '【参考文献】' in line_stripped or line_stripped.startswith('参考文献'):
            in_revised_contract = False
            in_suggestions = False
            continue
        
        if in_revised_contract:
            revised_contract += line + '\n'
            continue
        
        if in_suggestions:
            if line_stripped.startswith('1.') or line_stripped.startswith('2.') or \
               line_stripped.startswith('3.') or line_stripped.startswith('4.') or \
               line_stripped.startswith('5.') or line_stripped.startswith('6.') or \
               line_stripped.startswith('7.') or line_stripped.startswith('8.') or \
               line_stripped.startswith('9.') or line_stripped.startswith('10.'):
                if current_suggestion:
                    suggestions.append(current_suggestion)
                current_suggestion = {
                    'position': line_stripped.replace('修改位置：', '').strip(),
                    'original': '',
                    'revised': '',
                    'reason': ''
                }
            elif current_suggestion:
                if '原文：' in line_stripped:
                    current_suggestion['original'] = line_stripped.replace('原文：', '').strip()
                elif '修改为：' in line_stripped:
                    current_suggestion['revised'] = line_stripped.replace('修改为：', '').strip()
                elif '修改原因：' in line_stripped:
                    current_suggestion['reason'] = line_stripped.replace('修改原因：', '').strip()
                elif line_stripped:
                    if current_suggestion['reason']:
                        current_suggestion['reason'] += ' ' + line_stripped
                    elif current_suggestion['revised']:
                        current_suggestion['revised'] += ' ' + line_stripped
                    elif current_suggestion['original']:
                        current_suggestion['original'] += ' ' + line_stripped
    
    if current_suggestion:
        suggestions.append(current_suggestion)
    
    sections = []
    if revised_contract:
        sections.append({
            'name': '修订后的完整合同',
            'original': '（请查看修改建议了解具体修改内容）',
            'revised': revised_contract.strip(),
            'suggestions': []
        })
    
    return {
        'sections': sections,
        'suggestions': suggestions,
        'references': references,
        'revised_contract': revised_contract.strip()
    }


@app.route('/')
def index():
    """主页"""
    return render_template('index.html')


@app.route('/api/contracts', methods=['GET'])
def list_contracts():
    """列出可用的合同文件"""
    contracts_dir = 'data/contracts'
    
    if not os.path.exists(contracts_dir):
        return jsonify([])
    
    contracts = []
    for file in os.listdir(contracts_dir):
        if file.lower().endswith(('.docx', '.pdf')):
            contracts.append(file)
    
    return jsonify(sorted(contracts))


@app.route('/api/process', methods=['POST'])
def process_contract():
    """处理合同"""
    data = request.json
    filename = data.get('filename')
    
    if not filename:
        return jsonify({'error': '请提供文件名'}), 400
    
    contracts_dir = 'data/contracts'
    input_file_path = os.path.join(contracts_dir, filename)
    
    if not os.path.exists(input_file_path):
        return jsonify({'error': f'文件 {filename} 不存在'}), 404
    
    try:
        init_agent()
        
        output_dir = 'data/outputs'
        os.makedirs(output_dir, exist_ok=True)
        
        output_filename = f"{os.path.splitext(filename)[0]}(revised).txt"
        output_file_path = os.path.join(output_dir, output_filename)
        
        result = agent.process_contract(
            input_file_path,
            output_dir=output_dir,
            output_file=output_filename
        )
        
        if not result:
            return jsonify({'error': '处理失败'}), 500
        
        with open(output_file_path, 'r', encoding='utf-8') as f:
            result_text = f.read()
        
        parsed_data = parse_contract_result(result_text)
        
        return jsonify({
            'success': True,
            'filename': output_filename,
            'sections': parsed_data['sections'],
            'suggestions': parsed_data['suggestions'],
            'references': parsed_data['references'],
            'revised_contract': parsed_data['revised_contract']
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            'error': f'处理过程中发生错误: {str(e)}',
            'traceback': traceback.format_exc()
        }), 500


@app.route('/api/diff', methods=['POST'])
def compute_diff():
    """计算两个文本的差异"""
    data = request.json
    original = data.get('original', '')
    revised = data.get('revised', '')
    
    import difflib
    
    differ = difflib.Differ()
    diff = list(differ.compare(original.splitlines(), revised.splitlines()))
    
    diff_html = ''
    for line in diff:
        if line.startswith('  '):
            diff_html += f'<div class="diff-unchanged">{line[2:]}</div>'
        elif line.startswith('- '):
            diff_html += f'<div class="diff-deleted">{line[2:]}</div>'
        elif line.startswith('+ '):
            diff_html += f'<div class="diff-added">{line[2:]}</div>'
    
    return jsonify({'diff_html': diff_html})


@app.route('/api/download', methods=['POST'])
def download_word():
    """下载修订后的合同为Word文档"""
    data = request.json
    revised_contract = data.get('revised_contract', '')
    original_filename = data.get('filename', 'contract')
    
    if not revised_contract:
        return jsonify({'error': '没有可下载的修订合同内容'}), 400
    
    try:
        doc = Document()
        
        title = doc.add_heading('修订后的合同', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in title.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        paragraphs = revised_contract.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text)
                para.paragraph_format.line_spacing = 1.5
                
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        base_name = os.path.splitext(original_filename)[0]
        base_name = re.sub(r'\(revised\)$', '', base_name, flags=re.IGNORECASE)
        download_filename = f"{base_name}(revised).docx"
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=download_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        import traceback
        return jsonify({
            'error': f'生成Word文档失败: {str(e)}',
            'traceback': traceback.format_exc()
        }), 500


if __name__ == '__main__':
    print("=" * 80)
    print("合同修订智能体 - Web UI 服务")
    print("=" * 80)
    print("\n正在启动服务...")
    print("访问地址: http://localhost:5000")
    print("=" * 80)
    
    init_agent()
    
    app.run(debug=True, host='0.0.0.0', port=5000)
